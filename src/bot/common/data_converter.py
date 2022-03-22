from datetime import datetime, timedelta
from typing import List, Optional

import pymorphy2
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from common import calc_oof_profit, PayOffCalculation, calc_payoff_profit
from common.oof_profit_calculator import OOFCalculation
from repository import Repository


def convert_to_doc(data: dict) -> Document:
    claim_doc: Document = Document()

    # common doc settings
    section = claim_doc.sections[-1]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.6)

    # header
    header: Paragraph = claim_doc.add_paragraph()
    header.text = get_head_text(data["claim_data"]["head"])
    header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # theme
    theme: Paragraph = claim_doc.add_paragraph()
    theme.text = f"\nИсковое заявление\n{data['claim_theme']}"
    theme.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # story
    story_parts: List[str] = get_story_parts(data["claim_data"])
    for story_part in story_parts:
        story_paragraph: Paragraph = claim_doc.add_paragraph()
        story_paragraph.text = story_part
        story_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        story_paragraph.paragraph_format.first_line_indent = Inches(0.5)

    # essence
    essence: Paragraph = claim_doc.add_paragraph()
    essence_options: List[str] = data["claim_data"]["essence"]["chosen_options"]
    essence.add_run(", ".join(essence_options))
    essence.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    essence.paragraph_format.first_line_indent = Inches(0.5)

    # proofs
    proofs: Paragraph = claim_doc.add_paragraph()
    proofs_data: List[str] = data["claim_data"]["proofs"]["chosen_options"]
    proofs.text = "Доводы, указанные в исковом заявлении подтверждаются следующими " + \
                  f"доказательствами: {', '.join(proofs_data)}."
    proofs.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    proofs.paragraph_format.first_line_indent = Inches(0.5)

    # law
    law: Paragraph = claim_doc.add_paragraph()
    law.text = get_law_text(data["claim_theme"])
    law.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    law.paragraph_format.first_line_indent = Inches(0.5)

    # claims
    pre_claims: Paragraph = claim_doc.add_paragraph()
    pre_claims.text = "Прошу"
    pre_claims.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # don't use style="List Number" due to issue https://github.com/python-openxml/python-docx/issues/25
    claims: Paragraph = claim_doc.add_paragraph()
    for i, claim_data in enumerate(data["claim_data"]["claims"]["claims"], start=1):
        claims.add_run(f"{i}. {claim_data}\n")

    # additions
    pre_additions: Paragraph = claim_doc.add_paragraph()
    pre_additions.text = "Приложение"
    pre_additions.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    additions: Paragraph = claim_doc.add_paragraph()
    for i, addition_data in enumerate(data["claim_data"]["additions"]["chosen_options"], start=1):
        additions.add_run(f"{i}. {addition_data}\n")

    # footer
    footer: Paragraph = claim_doc.add_paragraph()
    footer.text = get_footer_text(data["claim_data"]["head"])
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    return claim_doc


def get_footer_text(head_data: dict) -> str:
    short_user_name: str = get_short_user_name(head_data['user_name'])
    claim_date: datetime = datetime.now()
    return f"{claim_date.strftime('%d.%m.%Y'): <50}{short_user_name}"


def get_head_text(head_data: dict) -> str:
    full_employer_name: str = get_full_employee_name(head_data)
    lines = [
        f"В {head_data['chosen_court'][0]}",
        f"Адрес: {head_data['chosen_court'][1]}",
        "",
        f"Истец: {head_data['user_name']}",
        f"Адрес: {head_data['user_post_code']}, г. {head_data['chosen_city']}, "
        f"ул. {head_data['chosen_street']}, д. {head_data['house_chosen']}"
        f"{', кв. ' + head_data['apartment_chosen'] if head_data['apartment_chosen'] != '' else ''}",
        "",
        f"Ответчик: {full_employer_name}",
        f"Адрес: {head_data['chosen_employer_address']}"
    ]

    return '\n'.join(lines)


def get_story_parts(data: dict) -> List[str]:
    full_employer_name: str = get_full_employee_name(data["head"])
    inflected_position_ablt: str = inflect(data["story"]["user_position"], "ablt")
    inflected_position_gent: str = inflect(data["story"]["user_position"], "gent")
    lines = [
        f"Я, {data['head']['user_name']}, работал {inflected_position_ablt} в {full_employer_name} "
        f"с {data['story']['start_work_date'].strftime('%d.%m.%Y')} на основании трудового договора, "
        f"в соответствии с которым был принят на работу к ответчику на должность {inflected_position_gent} "
        f"с окладом {data['story']['user_salary']} рублей.",
        data["story"]["story_conflict"]
    ]

    if data["story"]["user_employer_discussion"] != "":
        lines.append(data["story"]["user_employer_discussion"])
    return lines


def get_full_employee_name(head_data: dict) -> str:
    if head_data.get("inn") is None:
        return head_data["chosen_employer_name"]
    else:
        return f"{head_data['chosen_employer_name']} (ИНН {head_data['inn']})"


def get_law_text(claim_theme: str) -> str:
    repository: Repository = Repository()
    law_data: Optional[List[str]] = repository.get_claim_tmp_options(claim_theme, "law")
    if law_data is None:
        return ""

    return f"На основании изложенного, руководствуясь {', '.join(law_data)}."


def get_short_user_name(user_name: str) -> str:
    soname, name, second_name = tuple(user_name.split(" "))
    return f"{name[0].upper()}.{second_name[0].upper()}. {soname}"


def get_oof_profit_calculation(claim_data: dict) -> Document:
    end_work_date: datetime = claim_data["story"]["end_work_date"]
    start_oof_date: datetime = end_work_date + timedelta(days=1)
    avr_salary = claim_data["story"]["avr_salary"]
    oof_profit_calc: OOFCalculation = calc_oof_profit(start_oof_date, datetime.now(), avr_salary)

    calc_doc: Document = Document()

    # common doc settings
    section = calc_doc.sections[-1]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.6)

    theme: Paragraph = calc_doc.add_paragraph()
    theme_font = theme.add_run("Расчет задолженности по заработной плате за время вынужденного прогула").font

    theme.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    theme_font.size = Pt(14)
    theme_font.bold = True

    calc: Paragraph = calc_doc.add_paragraph()
    avr_salary_title_font = calc.add_run("Средняя заработная плата за предыдущий период\n").font
    avr_salary_title_font.bold = True
    calc.add_run("Среднее число рабочих дней в месяце: 20\n"
                 f"Cредняя заработная плата в месяц: {avr_salary}\n"
                 f"Средний заработок за день: {avr_salary} / 20 = {avr_salary/20}\n\n")

    day_oof_title_font = calc.add_run("Время вынужденного прогула\n").font
    day_oof_title_font.bold = True
    if oof_profit_calc.first_month_days_oof > 0 and oof_profit_calc.oof_months > 0:
        calc.add_run(
            f"Дата увольнения: {end_work_date.strftime('%d.%m.%Y')}\n"
            f"Дата начала вынужденного прогула: {start_oof_date.strftime('%d.%m.%Y')}\n"
            f"Дата подачи искового заявления: {datetime.now().strftime('%d.%m.%Y')}\n"
            f"Число рабочих дней за время первого месяца вынужденного прогула: {oof_profit_calc.first_month_days_oof}\n"
            f"Число полных месяцев за время вынужденного прогула: {oof_profit_calc.oof_months}\n"
            f"Число рабочих дней вынужденного прогула в текущем месяце: {oof_profit_calc.current_month_days_oof}\n"
            f"Число рабочих дней за время вынужденного прогула: {oof_profit_calc.oof_months} * 20 + "
            f"{oof_profit_calc.first_month_days_oof} + {oof_profit_calc.current_month_days_oof} = "
            f"{oof_profit_calc.oof_days}\n\n"
        )
    elif oof_profit_calc.first_month_days_oof > 0 and oof_profit_calc.oof_months == 0:
        calc.add_run(
            f"Дата увольнения: {end_work_date.strftime('%d.%m.%Y')}\n"
            f"Дата начала вынужденного прогула: {start_oof_date.strftime('%d.%m.%Y')}\n"
            f"Дата подачи искового заявления: {datetime.now().strftime('%d.%m.%Y')}\n"
            f"Число рабочих дней за время первого месяца вынужденного прогула: {oof_profit_calc.first_month_days_oof}\n"
            f"Число рабочих дней вынужденного прогула в текущем месяце: {oof_profit_calc.current_month_days_oof}\n"
            f"Число рабочих дней за время вынужденного прогула: {oof_profit_calc.first_month_days_oof} + "
            f"{oof_profit_calc.current_month_days_oof} = {oof_profit_calc.oof_days}\n\n"
        )
    elif oof_profit_calc.first_month_days_oof == 0:
        calc.add_run(
            f"Дата увольнения: {end_work_date.strftime('%d.%m.%Y')}\n"
            f"Дата начала вынужденного прогула: {start_oof_date.strftime('%d.%m.%Y')}\n"
            f"Дата подачи искового заявления: {datetime.now().strftime('%d.%m.%Y')}\n"
            f"Число рабочих дней вынужденного прогула в текущем месяце: {oof_profit_calc.current_month_days_oof}\n"
            f"Число рабочих дней за время вынужденного прогула: {oof_profit_calc.oof_days}\n\n"
        )

    oof_profit_title_font = calc.add_run("Сумма задолженности за время вынужденного прогула\n").font
    oof_profit_title_font.bold = True
    calc.add_run("{средний заработок за день} * {Число рабочих дней за время вынужденного прогула} = "
                 f"{oof_profit_calc.oof_days} * {avr_salary/20} = {oof_profit_calc.oof_profit}\n")

    footer: Paragraph = calc_doc.add_paragraph()
    footer.text = get_footer_text(claim_data["head"])
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    return calc_doc


def get_payoff_profit_calculation(claim_data: dict) -> Document:
    payoff_date: datetime = claim_data["story"]["payoff_date"]
    current_date: datetime = datetime.now()
    payoff_profit_calc: PayOffCalculation = calc_payoff_profit(payoff_date,
                                                               claim_data["story"]["pay_day_1"],
                                                               claim_data["story"]["payment_1"],
                                                               claim_data["story"]["pay_day_2"],
                                                               claim_data["story"]["payment_2"],
                                                               current_date)

    calc_doc: Document = Document()

    # common doc settings
    section = calc_doc.sections[-1]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(1)
    section.right_margin = Inches(0.6)

    theme: Paragraph = calc_doc.add_paragraph()
    theme_font = theme.add_run("Расчет задолженности по заработной плате").font

    theme.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    theme_font.size = Pt(14)
    theme_font.bold = True

    calc: Paragraph = calc_doc.add_paragraph()
    payoff_calc_title_font = calc.add_run("Заработная плата за предыдущий период\n").font
    payoff_calc_title_font.bold = True
    calc.add_run(
        f"Дата, когда перестала поступать заработная плата: {payoff_date.strftime('%d.%m.%Y')}\n"
        f"Число месяца, когда приходит заработная плата: {claim_data['story']['pay_day_1']}\n"
        f"Величина заработной платы, которая приходит {claim_data['story']['pay_day_1']}-го числа: {claim_data['story']['payment_1']}\n"
        f"Число просроченных платежей по заработной плате на текущую дату: {payoff_profit_calc.paydays_1_count}\n"
        f"Число месяца, когда приходит аванс: {claim_data['story']['pay_day_2']}\n"
        f"Величина аванса, которая приходит {claim_data['story']['pay_day_2']}-го числа: {claim_data['story']['payment_2']}\n"
        f"Число просроченных платежей по авансу на текущую дату: {payoff_profit_calc.paydays_2_count}\n"
        f"Общая сумма задолженности по заработной плате: {payoff_profit_calc.paydays_1_count} * {claim_data['story']['payment_1']} + "
        f"{payoff_profit_calc.paydays_2_count} * {claim_data['story']['payment_2']} = {payoff_profit_calc.payoff_profit}\n\n"
    )

    compensation_title_font = calc.add_run("Компенсация за предыдущий период\n").font
    compensation_title_font.bold = True

    calc.add_run(
        f"Число дней, прошедших с момента окончания выплат: {payoff_profit_calc.whole_days}\n"
        f"Ставка рефинансирования: {payoff_profit_calc.key_rate}%\n"
        f"Сумма компенсации: {payoff_profit_calc.payoff_profit} * (({payoff_profit_calc.key_rate} / 100) * 1/150) * "
        f"{payoff_profit_calc.whole_days} = {payoff_profit_calc.compensation}\n"
    )

    footer: Paragraph = calc_doc.add_paragraph()
    footer.text = get_footer_text(claim_data["head"])
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    return calc_doc


def inflect(word: str, case: str) -> str:
    morph = pymorphy2.MorphAnalyzer()
    word_analysis = morph.parse(word)[0]
    word_in_case = word_analysis.inflect({case})
    return word_in_case.word
